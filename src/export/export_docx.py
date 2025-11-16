from docx import Document
import io

def export_docx(data,df):
    doc=Document()
    idf=data["identification"]
    doc.add_heading("Báo cáo thẩm định",1)
    doc.add_paragraph(f"Họ tên: {idf['ten']}")
    doc.add_paragraph(f"CCCD: {idf['cccd']}")
    table=doc.add_table(rows=1,cols=5)
    h=table.rows[0].cells
    h[0].text="Tháng";h[1].text="TT";h[2].text="Lãi";h[3].text="Gốc";h[4].text="Dư"
    for _,r in df.head(10).iterrows():
        c=table.add_row().cells
        c[0].text=str(int(r["month"]))
        c[1].text=str(int(r["payment"]))
        c[2].text=str(int(r["interest"]))
        c[3].text=str(int(r["principal"]))
        c[4].text=str(int(r["balance"]))
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()
