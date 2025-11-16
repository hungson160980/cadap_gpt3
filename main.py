import streamlit as st
import pandas as pd
from docx import Document
import altair as alt
import io
import re

st.set_page_config(page_title="Thẩm định vay vốn", layout="wide")

# ======================================================
# 1) HÀM PARSER DOCX (TÍCH HỢP TRỰC TIẾP)
# ======================================================
def parse_vnd(s):
    if not s: return 0
    s = s.replace(".", "").replace(",", "").replace(" ", "")
    try: return int(s)
    except: return 0

def first(pattern, text):
    m = re.search(pattern, text, re.IGNORECASE)
    return m.group(1).strip() if m else None

def parse_docx(content):
    doc = Document(content)
    text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    lines = text.splitlines()

    # ==== Identification ====
    ten = first(r"Họ và tên[:\s]*(.+)", text) or ""
    cccd = first(r"(?:CCCD|CMND)[:\s]*([0-9]{9,12})", text) or ""
    diachi = first(r"(?:Địa chỉ|Nơi cư trú)[:\s]*(.+)", text) or ""
    phone = first(r"(?:Điện thoại|SĐT|Phone)[:\s]*([\d\+\-\s]{7,20})", text) or ""

    # ==== Finance ====
    mucdich = first(r"Mục đích[:\s]*(.+)", text) or ""
    tongnhucau = parse_vnd(first(r"Tổng nhu cầu.*?([\d\., ]+)", text) or "0")
    vondoing = parse_vnd(first(r"Vốn đối ứng.*?([\d\., ]+)", text) or "0")
    sotienvay = parse_vnd(first(r"(?:Số tiền vay|Vốn vay).*?([\d\., ]+)", text) or str(tongnhucau - vondoing))
    laisuat = float((first(r"Lãi suất[:\s]*([\d\.,]+)", text) or "8.5").replace(",", "."))
    thoihan = int(first(r"Thời hạn.*?(\d+)", text) or "60")

    # ==== Income ====
    thunhap = parse_vnd(first(r"Thu nhập.*?([\d\., ]+)", text) or "0")
    chiphi = parse_vnd(first(r"Chi phí.*?([\d\., ]+)", text) or "0")

    # ==== Collateral (simple detect) ====
    ts = []
    for i, L in enumerate(lines):
        if "giá trị" in L.lower() or "tài sản" in L.lower() or "bất động sản" in L.lower():
            val = parse_vnd(first(r"([\d\., ]+)\s*đ", L) or "0")
            addr = ""
            for n in range(i, min(i+5, len(lines))):
                if "địa chỉ" in lines[n].lower():
                    addr = re.sub(r"địa chỉ[:\s]*", "", lines[n], flags=re.I)
                    break
            ts.append({
                "loai": "TSĐB",
                "gia_tri": val,
                "dia_chi": addr,
                "ltv_percent": 0,
                "giay_to": ""
            })

    if not ts:
        ts.append({"loai":"TSĐB","gia_tri":0,"dia_chi":"","ltv_percent":0,"giay_to":""})

    return {
        "identification": {
            "ten": ten,
            "cccd": cccd,
            "dia_chi": diachi,
            "phone": phone
        },
        "finance": {
            "muc_dich": mucdich,
            "tong_nhu_cau": tongnhucau,
            "von_doi_ung": vondoing,
            "so_tien_vay": sotienvay,
            "lai_suat_p_a": laisuat,
            "thoi_han_thang": thoihan
        },
        "income": {
            "thu_nhap_hang_thang": thunhap,
            "chi_phi_hang_thang": chiphi
        },
        "collateral": ts
    }

# ======================================================
# 2) TÍNH TOÁN DÒNG TIỀN
# ======================================================
def monthly_payment(P, r, n):
    r = r/100/12
    if r == 0: return P/n
    return P * (r * (1 + r)**n) / ((1 + r)**n - 1)

def build_schedule(P, r, n):
    pay = monthly_payment(P, r, n)
    bal = P
    rows = []
    r_m = r/100/12
    for m in range(1, n+1):
        interest = bal * r_m
        principal = pay - interest
        bal -= principal
        rows.append({
            "month": m,
            "payment": round(pay),
            "interest": round(interest),
            "principal": round(principal),
            "balance": round(bal if bal > 0 else 0)
        })
    return pd.DataFrame(rows)

# ======================================================
# 3) EXPORT EXCEL
# ======================================================
def export_excel(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df.to_excel(w, index=False)
    buf.seek(0)
    return buf

# ======================================================
# 4) EXPORT DOCX
# ======================================================
def export_docx(full_data, df):
    doc = Document()
    idf = full_data["identification"]
    fin = full_data["finance"]

    doc.add_heading("BÁO CÁO THẨM ĐỊNH PHƯƠNG ÁN", 1)
    doc.add_paragraph(f"Họ tên: {idf['ten']}")
    doc.add_paragraph(f"CCCD: {idf['cccd']}")
    doc.add_paragraph(f"Địa chỉ: {idf['dia_chi']}")
    doc.add_paragraph(f"SĐT: {idf['phone']}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Mục đích vay: {fin['muc_dich']}")
    doc.add_paragraph(f"Số tiền vay: {fin['so_tien_vay']:,} đồng")
    doc.add_paragraph(f"Lãi suất: {fin['lai_suat_p_a']}%/năm")
    doc.add_paragraph(f"Thời hạn vay: {fin['thoi_han_thang']} tháng")
    doc.add_page_break()

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text="Tháng"; hdr[1].text="Thanh toán"; hdr[2].text="Lãi"; hdr[3].text="Gốc"; hdr[4].text="Dư nợ"

    for _, r in df.head(24).iterrows():
        row = table.add_row().cells
        row[0].text = str(r["month"])
        row[1].text = f"{r['payment']:,}"
        row[2].text = f"{r['interest']:,}"
        row[3].text = f"{r['principal']:,}"
        row[4].text = f"{r['balance']:,}"

    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# ======================================================
# 5) STREAMLIT APP
# ======================================================
st.title("📝 Thẩm định phương án sử dụng vốn (Bản 1 file đơn giản)")

uploaded = st.file_uploader("Tải lên file PASDV (.docx)", type=["docx"])

if "data" not in st.session_state:
    st.session_state.data = None

if uploaded:
    st.session_state.data = parse_docx(uploaded)
    st.success("Đọc file thành công!")

# default template
if st.session_state.data is None:
    st.session_state.data = {
        "identification":{"ten":"","cccd":"","dia_chi":"","phone":""},
        "finance":{"muc_dich":"","tong_nhu_cau":0,"von_doi_ung":0,"so_tien_vay":0,"lai_suat_p_a":8.5,"thoi_han_thang":60},
        "income":{"thu_nhap_hang_thang":0,"chi_phi_hang_thang":0},
        "collateral":[{"loai":"TSĐB","gia_tri":0,"dia_chi":"","ltv_percent":0}]
    }

data = st.session_state.data

# ===== SHOW BASIC INFO =====
idf = data["identification"]
fin = data["finance"]

st.subheader("🔎 Thông tin khách hàng")
st.write(idf)

st.subheader("💰 Thông tin phương án")
st.write(fin)

# ====== CALC ======
df = build_schedule(fin["so_tien_vay"], fin["lai_suat_p_a"], fin["thoi_han_thang"])

st.subheader("📊 Lịch trả nợ (24 tháng đầu)")
st.dataframe(df.head(24))

# Chart
chart = alt.Chart(df).mark_line().encode(
    x="month",
    y="payment"
).properties(width=800)
st.altair_chart(chart, use_container_width=True)

# ===== EXPORT =====
col1, col2 = st.columns(2)
with col1:
    st.subheader("📤 Xuất Excel")
    st.download_button("Tải file Excel", export_excel(df), file_name="ke_hoach.xlsx")

with col2:
    st.subheader("📤 Xuất DOCX")
    st.download_button("Tải file DOCX", export_docx(data, df), file_name="bao_cao.docx")
